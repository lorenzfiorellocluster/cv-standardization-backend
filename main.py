import os
from jinja2 import Environment, FileSystemLoader
import re

# -------------------------------------------------------
# LATEX ESCAPE FILTER (FONDAMENTALE!)
# -------------------------------------------------------
def latex_escape(text):
    """Escapa caratteri speciali LaTeX"""
    if text is None:
        return ""
    
    text = str(text)
    
    # Caratteri speciali LaTeX da escapare
    # ORDINE IMPORTANTE: \ per primo, poi gli altri
    replacements = [
        ('\\', r'\textbackslash '),  # backslash per primo
        ('&', r'\&'),
        ('%', r'\%'),
        ('$', r'\$'),
        ('#', r'\#'),
        ('_', r'\_'),
        ('{', r'\{'),
        ('}', r'\}'),
        ('~', r'\textasciitilde '),
        ('^', r'\^{}'),
    ]
    
    for old, new in replacements:
        text = text.replace(old, new)
    
    return text
# -------------------------------------------------------
# JSON DI TEST
# -------------------------------------------------------
cv = {
    "id": "00001",
    "personal_info": {
        "full_name": "John Doe",
        "role": "Cloud Engineering Consultant",
        "department": "Consultant – Cloud Engineering",
        "seniority": "Senior",
        "location": "Berlin, Germany",
        "email": "john.doe@example.com",
        "phone": "+49 151 12345678"
    },

    "summary": "Cloud Engineering Consultant with hands-on experience on Azure, AWS, Terraform and DevOps automation.",
    "key_achievements": [
        "30% faster deployments through CI/CD optimization",
        "35% reduction in code review time via AI-powered tooling",
        "30% cost optimization on cloud workloads",
        "90% faster environment provisioning with Terraform modules"
    ],

    "key_skills": [
        {
            "category": "Cloud Platforms",
            "items": [
                "Azure (2 yrs): AKS, App Services, Storage, VNets",
                "AWS (1.5 yrs): EC2, Lambda, RDS, CloudWatch"
            ]
        },
        {
            "category": "Programming & Scripting",
            "items": [
                "PowerShell (2 yrs)",
                "Python",
                "Bash",
                "YAML",
                "JSON"
            ]
        },
        {
            "category": "Infrastructure as Code",
            "items": [
                "Terraform (2 yrs)",
                "Bicep",
                "CloudFormation"
            ]
        },
        {
            "category": "Containers & Orchestration",
            "items": [
                "Docker (2 yrs)",
                "Kubernetes/AKS",
                "Azure Container Apps"
            ]
        }
    ],

    "certifications": [
        { "name": "Azure Administrator", "code": "AZ-104", "year": "2025" },
        { "name": "Terraform Associate", "code": "003", "year": "2025" },
        { "name": "Solutions Architect", "code": "SAA-C03", "year": "2024" }
    ],

    "languages": [
        { "language": "German", "level": "Native" },
        { "language": "English", "level": "Fluent (C2)" }
    ],

    "experience": [
        {
            "role": "Cloud Engineer",
            "company": "Cluster Reply GmbH",
            "start_date": "April 2025",
            "end_date": "Present",
            "description_list": [
                "Cloud engineering, IaC automation, DevOps",
                "Speaker at Developer Summit Berlin 2025: AI Security & Workflows"
            ]
        }
    ]
}

# -------------------------------------------------------
# PATHS
# -------------------------------------------------------
TEMPLATE_DIR = "template"
OUTPUT_DIR = "output"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# LOGO PATH FIX (absolute path)
cv["logo_path"] = os.path.abspath("assets/cluster_reply_logo.jpg")

# -------------------------------------------------------
# JINJA RENDERING CON FILTRO CUSTOM
# -------------------------------------------------------
env = Environment(
    loader=FileSystemLoader(TEMPLATE_DIR),
    block_start_string="<%",
    block_end_string="%>",
    variable_start_string="<<",
    variable_end_string=">>",
    comment_start_string="<#",
    comment_end_string="#>"
)

# AGGIUNGI IL FILTRO 'tex'
env.filters['tex'] = latex_escape

template = env.get_template("cv_template.tex")
latex_output = template.render(**cv)  # Nota: **cv invece di cv=cv

tex_path = os.path.join(OUTPUT_DIR, "cv_output.tex")

with open(tex_path, "w", encoding="utf-8") as f:
    f.write(latex_output)

print("✔ Generated LaTeX:", tex_path)

# -------------------------------------------------------
# DOCX
# -------------------------------------------------------
from docx import Document

doc = Document()
doc.add_heading(cv["personal_info"]["full_name"], level=1)
doc.add_paragraph(cv["summary"])

doc.add_heading("Key Achievements", level=2)
for ach in cv["key_achievements"]:
    doc.add_paragraph(f"- {ach}")

doc_path = os.path.join(OUTPUT_DIR, "cv_output.docx")
doc.save(doc_path)

print("✔ Generated DOCX:", doc_path)